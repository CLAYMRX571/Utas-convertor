import time
from pathlib import Path
import pdfplumber
import requests
from deep_translator import GoogleTranslator
from docx import Document
from telebot import types
from keys import (
    ADMIN_PHONE,
    ADMIN_USERNAME,
    ILOVEPDF_PUBLIC_KEY,
    ILOVEPDF_SECRET_KEY,
    LANG_MAP,
)
from states import (
    DOCX_DIR,
    OUT_DIR,
    PDF_DIR,
    clear_state,
    get_state,
    set_state,
    unique_path,
)

def build_main_menu():
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.row("Pdf 📁", "Word 📁")
    markup.row("Text ✉️", "Tarjima 🌐")
    markup.row("Adminga murojaat 👨‍💻")
    return markup

def build_confirm_menu():
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.row("Ha ✅", "Yo'q ❌")
    markup.row("Orqaga 🔙")
    return markup

def build_translate_menu():
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    buttons = list(LANG_MAP.keys())
    for i in range(0, len(buttons), 2):
        markup.row(*buttons[i:i + 2])
    markup.row("Orqaga 🔙")
    return markup

MAIN_MENU = build_main_menu()
CONFIRM_MENU = build_confirm_menu()
TRANSLATE_MENU = build_translate_menu()

def send_progress(bot, chat_id: int, title: str):
    msg = bot.send_message(chat_id, f"{title}\n0%")
    for percent in range(10, 101, 10):
        time.sleep(0.15)
        try:
            bot.edit_message_text(
                chat_id=chat_id,
                message_id=msg.message_id,
                text=f"{title}\n{percent}%"
            )
        except Exception:
            pass
    return msg


def download_telegram_file(bot, file_id: str, dest: Path) -> Path:
    tg_file = bot.get_file(file_id)
    downloaded_file = bot.download_file(tg_file.file_path)
    dest.write_bytes(downloaded_file)
    return dest


def ilovepdf_auth() -> str:
    if not ILOVEPDF_PUBLIC_KEY or not ILOVEPDF_SECRET_KEY:
        raise RuntimeError("ILOVEPDF kalitlari topilmadi")

    response = requests.post(
        "https://api.ilovepdf.com/v1/auth",
        json={
            "public_key": ILOVEPDF_PUBLIC_KEY,
            "secret_key": ILOVEPDF_SECRET_KEY,
        },
        timeout=60,
    )
    response.raise_for_status()
    data = response.json()
    token = data.get("token")
    if not token:
        raise RuntimeError(f"iLovePDF auth xatolik: {data}")
    return token


def ilovepdf_start(tool: str, headers: dict) -> tuple[str, str]:
    response = requests.get(
        f"https://api.ilovepdf.com/v1/start/{tool}",
        headers=headers,
        timeout=60,
    )
    response.raise_for_status()
    data = response.json()

    server = data.get("server")
    task = data.get("task")

    if not server or not task:
        raise RuntimeError(f"iLovePDF start xatolik: {data}")

    return server, task


def ilovepdf_upload(server: str, task: str, headers: dict, file_path: Path) -> str:
    with file_path.open("rb") as f:
        response = requests.post(
            f"https://{server}/v1/upload",
            headers=headers,
            data={"task": task},
            files={"file": (file_path.name, f)},
            timeout=180,
        )
    response.raise_for_status()
    data = response.json()

    server_filename = data.get("server_filename")
    if not server_filename:
        raise RuntimeError(f"Upload xatolik: {data}")

    return server_filename


def ilovepdf_download(server: str, task: str, headers: dict, output_path: Path) -> Path:
    response = requests.get(
        f"https://{server}/v1/download/{task}",
        headers=headers,
        timeout=180,
    )
    response.raise_for_status()
    output_path.write_bytes(response.content)

    if not output_path.exists() or output_path.stat().st_size == 0:
        raise RuntimeError("Natija fayl bo'sh chiqdi")

    return output_path


def pdf_to_word_convert(pdf_path: Path) -> Path:
    token = ilovepdf_auth()
    headers = {"Authorization": f"Bearer {token}"}
    server, task = ilovepdf_start("pdfword", headers)
    server_filename = ilovepdf_upload(server, task, headers, pdf_path)

    process_response = requests.post(
        f"https://{server}/v1/process",
        headers=headers,
        json={
            "task": task,
            "tool": "pdfword",
            "files": [
                {
                    "server_filename": server_filename,
                    "filename": pdf_path.name,
                }
            ],
        },
        timeout=180,
    )
    process_response.raise_for_status()

    output_path = unique_path(OUT_DIR, "docx")
    return ilovepdf_download(server, task, headers, output_path)


def word_to_pdf_convert(docx_path: Path) -> Path:
    token = ilovepdf_auth()
    headers = {"Authorization": f"Bearer {token}"}
    server, task = ilovepdf_start("officepdf", headers)
    server_filename = ilovepdf_upload(server, task, headers, docx_path)

    process_response = requests.post(
        f"https://{server}/v1/process",
        headers=headers,
        json={
            "task": task,
            "tool": "officepdf",
            "files": [
                {
                    "server_filename": server_filename,
                    "filename": docx_path.name,
                }
            ],
        },
        timeout=180,
    )
    process_response.raise_for_status()

    output_path = unique_path(OUT_DIR, "pdf")
    return ilovepdf_download(server, task, headers, output_path)


def extract_text_from_docx(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)
    return "\n".join(parts).strip()


def extract_text_from_pdf(pdf_path: Path) -> str:
    texts = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text()
            if txt and txt.strip():
                texts.append(txt.strip())
    return "\n\n".join(texts).strip()


def translate_big_text(text: str, target_lang: str) -> str:
    if not text.strip():
        return ""

    translator = GoogleTranslator(source="auto", target=target_lang)
    chunks = []
    step = 3000

    for i in range(0, len(text), step):
        part = text[i:i + step]
        translated = translator.translate(part)
        if translated:
            chunks.append(translated)

    return "\n".join(chunks).strip()


def text_to_word_file(text: str) -> Path:
    output_path = unique_path(OUT_DIR, "docx")
    doc = Document()
    doc.add_heading("Foydalanuvchi matni", level=1)
    doc.add_paragraph(text)
    doc.save(str(output_path))
    return output_path


def send_file_and_remove(bot, chat_id: int, path: Path, caption: str):
    try:
        with open(path, "rb") as f:
            bot.send_document(
                chat_id,
                f,
                visible_file_name=path.name,
                caption=caption,
                reply_markup=MAIN_MENU,
            )
    finally:
        try:
            if path.exists():
                path.unlink()
        except Exception:
            pass


def register_handlers(bot):
    @bot.message_handler(commands=["start"])
    def start_handler(message):
        clear_state(message.from_user.id)
        username = message.from_user.full_name if message.from_user else "Foydalanuvchi"
        bot.send_message(
            message.chat.id,
            f"Assalomu alaykum, {username}\n\nKerakli bo'limni tanlang 👇",
            reply_markup=MAIN_MENU,
        )

    @bot.message_handler(func=lambda m: m.text == "Orqaga 🔙")
    def back_handler(message):
        clear_state(message.from_user.id)
        bot.send_message(message.chat.id, "Orqaga qaytdi 🔙", reply_markup=MAIN_MENU)

    @bot.message_handler(func=lambda m: m.text == "Pdf 📁")
    def pdf_start(message):
        clear_state(message.from_user.id)
        set_state(message.from_user.id, step="waiting_pdf")
        bot.send_message(message.chat.id, "PDF file tashlang 📁", reply_markup=MAIN_MENU)

    @bot.message_handler(func=lambda m: m.text == "Word 📁")
    def word_start(message):
        clear_state(message.from_user.id)
        set_state(message.from_user.id, step="waiting_word")
        bot.send_message(message.chat.id, "Word file tashlang (.docx) 📁", reply_markup=MAIN_MENU)

    @bot.message_handler(func=lambda m: m.text == "Text ✉️")
    def text_start(message):
        clear_state(message.from_user.id)
        set_state(message.from_user.id, step="waiting_text")
        bot.send_message(message.chat.id, "Matn yuboring ✍️", reply_markup=MAIN_MENU)

    @bot.message_handler(func=lambda m: m.text == "Tarjima 🌐")
    def translate_start(message):
        clear_state(message.from_user.id)
        set_state(message.from_user.id, step="waiting_translate_file")
        bot.send_message(message.chat.id, "Word yoki PDF formatida tashlang 📁", reply_markup=MAIN_MENU)

    @bot.message_handler(func=lambda m: m.text == "Adminga murojaat 👨‍💻")
    def admin_contact(message):
        text = (
            "👨‍💻 Admin bilan bog'lanish:\n\n"
            f"📩 Lichka: @{ADMIN_USERNAME}\n"
            f"📞 Telefon: {ADMIN_PHONE}"
        )
        bot.send_message(message.chat.id, text, reply_markup=MAIN_MENU)

    @bot.message_handler(content_types=["document"])
    def document_handler(message):
        data = get_state(message.from_user.id)
        step = data.get("step")
        document = message.document

        if not document or not document.file_name:
            bot.send_message(message.chat.id, "Fayl topilmadi ❌", reply_markup=MAIN_MENU)
            return

        filename = document.file_name.lower()

        if step == "waiting_pdf":
            if not filename.endswith(".pdf"):
                bot.send_message(message.chat.id, "Faqat PDF file tashlang 📁")
                return

            pdf_path = unique_path(PDF_DIR, "pdf")
            download_telegram_file(bot, document.file_id, pdf_path)
            set_state(
                message.from_user.id,
                step="confirm_pdf_to_word",
                file_path=str(pdf_path),
                file_type="pdf",
            )
            bot.send_message(message.chat.id, "Word ga o'girmoqchimisiz?", reply_markup=CONFIRM_MENU)
            return

        if step == "waiting_word":
            if not filename.endswith(".docx"):
                bot.send_message(message.chat.id, "Faqat .docx formatdagi Word file tashlang 📁")
                return

            docx_path = unique_path(DOCX_DIR, "docx")
            download_telegram_file(bot, document.file_id, docx_path)
            set_state(
                message.from_user.id,
                step="confirm_word_to_pdf",
                file_path=str(docx_path),
                file_type="docx",
            )
            bot.send_message(message.chat.id, "PDF ga o'girmoqchimisiz?", reply_markup=CONFIRM_MENU)
            return

        if step == "waiting_translate_file":
            if filename.endswith(".pdf"):
                file_path = unique_path(PDF_DIR, "pdf")
                file_type = "pdf"
            elif filename.endswith(".docx"):
                file_path = unique_path(DOCX_DIR, "docx")
                file_type = "docx"
            else:
                bot.send_message(message.chat.id, "Faqat PDF yoki DOCX file tashlang 📁")
                return

            download_telegram_file(bot, document.file_id, file_path)
            set_state(
                message.from_user.id,
                step="waiting_translate_lang",
                file_path=str(file_path),
                file_type=file_type,
            )
            bot.send_message(message.chat.id, "Tilni tanlang 👇", reply_markup=TRANSLATE_MENU)
            return

        bot.send_message(message.chat.id, "Avval kerakli menyuni tanlang 👇", reply_markup=MAIN_MENU)

    @bot.message_handler(func=lambda m: m.text == "Ha ✅")
    def confirm_yes(message):
        data = get_state(message.from_user.id)
        step = data.get("step")
        file_path = data.get("file_path")

        if not file_path:
            bot.send_message(message.chat.id, "Jarayon topilmadi. Qaytadan boshlang.", reply_markup=MAIN_MENU)
            clear_state(message.from_user.id)
            return

        src = Path(file_path)
        out = None

        try:
            if step == "confirm_pdf_to_word":
                prog = send_progress(bot, message.chat.id, "PDF Word ga o'girilmoqda 🔄")
                out = pdf_to_word_convert(src)
                try:
                    bot.edit_message_text("Tayyor ✅", message.chat.id, prog.message_id)
                except Exception:
                    pass
                send_file_and_remove(bot, message.chat.id, out, "Mana Word faylingiz 📁")
                clear_state(message.from_user.id)
                return

            if step == "confirm_word_to_pdf":
                prog = send_progress(bot, message.chat.id, "Word PDF ga o'girilmoqda 🔄")
                out = word_to_pdf_convert(src)
                try:
                    bot.edit_message_text("Tayyor ✅", message.chat.id, prog.message_id)
                except Exception:
                    pass
                send_file_and_remove(bot, message.chat.id, out, "Mana PDF faylingiz 📁")
                clear_state(message.from_user.id)
                return

            bot.send_message(message.chat.id, "Tasdiqlash uchun jarayon yo'q.", reply_markup=MAIN_MENU)

        except Exception as e:
            bot.send_message(message.chat.id, f"Xatolik yuz berdi: {e} ❌", reply_markup=MAIN_MENU)
            clear_state(message.from_user.id)
        finally:
            try:
                if src.exists():
                    src.unlink()
            except Exception:
                pass

    @bot.message_handler(func=lambda m: m.text == "Yo'q ❌")
    def confirm_no(message):
        clear_state(message.from_user.id)
        bot.send_message(message.chat.id, "Orqaga qaytdi 🔙", reply_markup=MAIN_MENU)

    @bot.message_handler(func=lambda m: m.text in list(LANG_MAP.keys()))
    def translate_lang(message):
        data = get_state(message.from_user.id)
        if data.get("step") != "waiting_translate_lang":
            bot.send_message(message.chat.id, "Avval fayl yuboring 📁", reply_markup=MAIN_MENU)
            return

        file_path = Path(data["file_path"])
        file_type = data["file_type"]
        lang_code = LANG_MAP[message.text]

        try:
            prog = send_progress(bot, message.chat.id, f"{lang_code} tiliga tarjima qilinmoqda 🔄")

            if file_type == "pdf":
                original_text = extract_text_from_pdf(file_path)
            else:
                original_text = extract_text_from_docx(file_path)

            if not original_text.strip():
                try:
                    bot.edit_message_text("File ichidan text topilmadi ❌", message.chat.id, prog.message_id)
                except Exception:
                    pass
                bot.send_message(
                    message.chat.id,
                    "File ichida tarjima qilinadigan text yo'q ❌",
                    reply_markup=MAIN_MENU,
                )
                clear_state(message.from_user.id)
                return

            translated_text = translate_big_text(original_text, lang_code)

            try:
                bot.edit_message_text("Tayyor ✅", message.chat.id, prog.message_id)
            except Exception:
                pass

            if not translated_text.strip():
                bot.send_message(message.chat.id, "Tarjima qilinmadi ❌", reply_markup=MAIN_MENU)
                clear_state(message.from_user.id)
                return

            parts = [translated_text[i:i + 4000] for i in range(0, len(translated_text), 4000)]
            for index, part in enumerate(parts):
                if index == len(parts) - 1:
                    bot.send_message(message.chat.id, part, reply_markup=MAIN_MENU)
                else:
                    bot.send_message(message.chat.id, part)

        except Exception as e:
            bot.send_message(message.chat.id, f"Xatolik yuz berdi: {e} ❌", reply_markup=MAIN_MENU)
        finally:
            clear_state(message.from_user.id)

    @bot.message_handler(content_types=["text"])
    def text_handler(message):
        data = get_state(message.from_user.id)

        if data.get("step") == "waiting_text":
            text = (message.text or "").strip()
            if not text:
                bot.send_message(message.chat.id, "Matn bo'sh bo'lmasligi kerak")
                return

            try:
                out = text_to_word_file(text)
                send_file_and_remove(bot, message.chat.id, out, "Word faylingiz tayyor ✅")
            except Exception as e:
                bot.send_message(message.chat.id, f"Xatolik yuz berdi: {e} ❌", reply_markup=MAIN_MENU)
            finally:
                clear_state(message.from_user.id)
            return

        bot.send_message(message.chat.id, "Kerakli tugmani tanlang 👇", reply_markup=MAIN_MENU)