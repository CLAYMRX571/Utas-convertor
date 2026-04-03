import os
import uuid
import re
from pathlib import Path
import fitz
import pdfplumber
import easyocr
import telebot
import requests
import subprocess
import shutil
import cv2
import pytesseract
from PyPDF2 import PdfReader
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from telebot.types import Message
from keys import get_menu, get_confirm, get_lang
from states import user_states, user_data, State

BASE_DIR = Path(__file__).resolve().parent
MEDIA_DIR = BASE_DIR / "media"
PDF_DIR = MEDIA_DIR / "pdf"
DOCX_DIR = MEDIA_DIR / "docx"
TRANSLATE_DIR = MEDIA_DIR / "translate"
OUT_DIR = MEDIA_DIR / "out"
TEMP_DIR = MEDIA_DIR / "temp"

for folder in [MEDIA_DIR, PDF_DIR, DOCX_DIR, TRANSLATE_DIR, OUT_DIR, TEMP_DIR]:
    folder.mkdir(parents=True, exist_ok=True)

def extract_text_from_docx_file_full(docx_path: str) -> str:
    try:
        doc = Document(docx_path)
        lines = []

        for para in doc.paragraphs:
            text = para.text
            if text is not None:
                text = normalize_ocr_chars(text)
                text = text.rstrip()
                if text != "":
                    lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text_parts = []
                    for para in cell.paragraphs:
                        t = para.text
                        if t is not None:
                            t = normalize_ocr_chars(t).strip()
                            if t:
                                cell_text_parts.append(t)
                    row_cells.append(" | ".join(cell_text_parts).strip())
                row_text = " ; ".join([c for c in row_cells if c.strip()])
                if row_text:
                    lines.append(row_text)

        return "\n".join(lines).strip()
    except Exception:
        return ""

def make_file_path(folder, suffix):
    folder = Path(folder)
    folder.mkdir(parents=True, exist_ok=True)
    return folder / f"{uuid.uuid4().hex}{suffix}"

def normalize_ocr_chars(text: str) -> str:
    if not text:
        return ""

    replacements = {
        '\x00': ' ',
        '|': ' ',
        '¦': ' ',
        '—': '-',
        '–': '-',
        '“': '"',
        '”': '"',
        '‘': "'",
        '’': "'",
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    return text

def has_too_many_repeated_chars(line: str) -> bool:
    return bool(re.search(r'(.)\1{4,}', line))

def is_garbage_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return True

    if len(stripped) <= 2:
        return True

    letters = sum(ch.isalpha() for ch in stripped)
    digits = sum(ch.isdigit() for ch in stripped)
    punct = sum(not ch.isalnum() and not ch.isspace() for ch in stripped)
    total = len(stripped)

    if total == 0:
        return True

    if letters == 0 and (digits > 0 or punct > 0):
        return True

    if total >= 6 and (digits + punct) / total > 0.60 and letters / total < 0.30:
        return True

    if has_too_many_repeated_chars(stripped):
        return True

    return False

def looks_like_ocr_noise(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return True

    tokens = stripped.split()
    total = len(stripped)
    letters = sum(ch.isalpha() for ch in stripped)
    digits = sum(ch.isdigit() for ch in stripped)
    punct = sum(not ch.isalnum() and not ch.isspace() for ch in stripped)

    if total == 0:
        return True

    if punct / total > 0.28 and letters / total < 0.45:
        return True

    if (digits + punct) / total > 0.55 and letters / total < 0.35:
        return True

    if len(tokens) >= 8:
        short_tokens = sum(1 for t in tokens if len(t) <= 2)
        if short_tokens / len(tokens) > 0.70 and letters / total < 0.50:
            return True

    return False

def is_meaningful_text(text: str) -> bool:
    if not text:
        return False
    text = text.strip()
    if len(text) < 20:
        return False
    alnum_count = sum(ch.isalnum() for ch in text)
    return alnum_count >= 10

def set_doc_defaults(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

def remove_file(file_path):
    try:
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
    except:
        pass

def extract_text_with_pdfplumber(pdf_path: str) -> str:
    texts = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    texts.append(page_text)
    except:
        pass
    return "\n\n".join(texts).strip()

def extract_text_with_pymupdf(pdf_path: str) -> str:
    texts = []

    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            page_text = page.get_text("text") or ""
            if page_text.strip():
                texts.append(page_text)
        doc.close()
    except:
        pass
    return "\n\n".join(texts).strip()

def extract_text_from_image(image_path):
    texts = []
    seen = set()

    readers = [
        easyocr.Reader(['en', 'tr'], gpu=False, verbose=False),
        easyocr.Reader(['ru', 'en'], gpu=False, verbose=False)
    ]

    for reader in readers:
        results = reader.readtext(image_path, detail=1)

        for bbox, text, prob in results:
            text = normalize_ocr_chars(text).strip()
            text = re.sub(r'\s+', ' ', text)

            if not text:
                continue

            if prob < 0.3:
                continue

            key = text.lower()
            if key in seen:
                continue

            seen.add(key)
            texts.append(text)

    return "\n".join(texts)

def extract_text_with_ocr(pdf_path: str) -> str:
    result_pages = []
    doc = fitz.open(pdf_path)

    try:
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)

            # OCR sifatini oshirish
            matrix = fitz.Matrix(2.5, 2.5)
            pix = page.get_pixmap(matrix=matrix, alpha=False)

            img_path = make_file_path(TEMP_DIR, ".png")
            pix.save(str(img_path))

            try:
                page_text = extract_text_from_image(str(img_path))
                if page_text and page_text.strip():
                    result_pages.append(page_text)
            finally:
                remove_file(str(img_path))
    finally:
        doc.close()

    return "\n\n".join(result_pages).strip()

def extract_text_from_pdf_full(pdf_path: str) -> str:
    text = extract_text_with_pdfplumber(pdf_path)
    if text and text.strip():
        return normalize_ocr_chars(text).strip()

    text = extract_text_with_pymupdf(pdf_path)
    if text and text.strip():
        return normalize_ocr_chars(text).strip()

    text = extract_text_with_ocr(pdf_path)
    if text and text.strip():
        return normalize_ocr_chars(text).strip()

    return ""

def create_docx_from_text(text: str, output_docx: str):
    doc = Document()
    set_doc_defaults(doc)

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        if is_garbage_line(line):
            continue

        if looks_like_ocr_noise(line):
            continue

        add_paragraph(doc, line)

    doc.save(output_docx)

def convert_docx_to_pdf_ilovepdf(docx_path: str, output_dir: str) -> str:
    output_dir = str(Path(output_dir).resolve())
    docx_path = str(Path(docx_path).resolve())

    if not os.path.exists(docx_path):
        raise Exception("DOCX fayl topilmadi.")

    token = ilovepdf_get_token()
    server, task = ilovepdf_start_task(token, tool="officepdf", region="eu")
    uploaded_file = ilovepdf_upload_file(token, server, task, docx_path)
    ilovepdf_process_officepdf(token, server, task, uploaded_file)

    output_pdf_path = str(Path(output_dir) / f"{Path(docx_path).stem}.pdf")
    ilovepdf_download_file(token, server, task, output_pdf_path)

    return output_pdf_path

def ilovepdf_get_token() -> str:
    public_key = os.getenv("ILOVEPDF_PUBLIC_KEY", "").strip()

    if not public_key:
        raise Exception(
            "ILOVEPDF_PUBLIC_KEY topilmadi. .env faylga yoki hosting environment ga qo‘ying."
        )

    url = "https://api.ilovepdf.com/v1/auth"
    resp = requests.post(url, data={"public_key": public_key}, timeout=60)

    if resp.status_code != 200:
        raise Exception(f"iLovePDF auth xatoligi: {resp.text}")

    data = resp.json()
    token = data.get("token")
    if not token:
        raise Exception("iLovePDF token olinmadi.")

    return token

def ilovepdf_start_task(token: str, tool: str = "officepdf", region: str = "eu") -> tuple[str, str]:
    url = f"https://api.ilovepdf.com/v1/start/{tool}/{region}"
    headers = {"Authorization": f"Bearer {token}"}

    resp = requests.get(url, headers=headers, timeout=60)
    if resp.status_code != 200:
        raise Exception(f"Start task xatoligi: {resp.text}")

    data = resp.json()
    server = data.get("server")
    task = data.get("task")

    if not server or not task:
        raise Exception("Server yoki task olinmadi.")

    return server, task

def ilovepdf_upload_file(token: str, server: str, task: str, file_path: str) -> dict:
    url = f"https://{server}/v1/upload"
    headers = {"Authorization": f"Bearer {token}"}

    with open(file_path, "rb") as f:
        files = {
            "file": (Path(file_path).name, f)
        }
        data = {
            "task": task
        }
        resp = requests.post(url, headers=headers, data=data, files=files, timeout=300)

    if resp.status_code != 200:
        raise Exception(f"Upload xatoligi: {resp.text}")

    uploaded = resp.json()
    if not uploaded.get("server_filename"):
        raise Exception("server_filename olinmadi.")

    return uploaded

def ilovepdf_process_officepdf(token: str, server: str, task: str, uploaded_file: dict) -> dict:
    url = f"https://{server}/v1/process"
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }

    payload = {
        "task": task,
        "tool": "officepdf",
        "files": [
            {
                "server_filename": uploaded_file["server_filename"],
                "filename": uploaded_file.get("filename", Path(uploaded_file["server_filename"]).name),
                "rotate": 0
            }
        ]
    }

    resp = requests.post(url, headers=headers, json=payload, timeout=300)
    if resp.status_code != 200:
        raise Exception(f"Process xatoligi: {resp.text}")

    return resp.json()

def ilovepdf_download_file(token: str, server: str, task: str, output_pdf_path: str):
    url = f"https://{server}/v1/download/{task}"
    headers = {"Authorization": f"Bearer {token}"}

    with requests.get(url, headers=headers, stream=True, timeout=300) as resp:
        if resp.status_code != 200:
            raise Exception(f"Download xatoligi: {resp.text}")

        with open(output_pdf_path, "wb") as f:
            for chunk in resp.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)

    if not os.path.exists(output_pdf_path) or os.path.getsize(output_pdf_path) == 0:
        raise Exception("PDF fayl yuklab olinmadi.")
    
def translate_text(text, target_lang):
    return GoogleTranslator(source='auto', target=target_lang).translate(text)

def smart_translate_full(text: str, target_lang: str) -> str:
    chunks = split_text_preserve_lines(text, max_length=3500)
    translated_chunks = []

    for chunk in chunks:
        chunk = chunk.strip("\n")
        if not chunk.strip():
            continue

        translated = GoogleTranslator(source='auto', target=target_lang).translate(chunk)
        if translated is None:
            translated = ""

        translated_chunks.append(str(translated))

    return "\n".join(translated_chunks).strip()

def split_text_preserve_lines(text: str, max_length: int = 3500):
    if not text:
        return []

    lines = text.splitlines()
    chunks = []
    current = ""

    for line in lines:
        if current:
            candidate = current + "\n" + line
        else:
            candidate = line

        if len(candidate) <= max_length:
            current = candidate
        else:
            if current:
                chunks.append(current)
            current = line

    if current:
        chunks.append(current)

    return chunks

def set_state(user_id, state):
    user_states[user_id] = state

def get_state(user_id):
    return user_states.get(user_id, State.NONE)

def clear_state(user_id):
    user_states[user_id] = State.NONE
    user_data.pop(user_id, None)

def register_handlers(bot: telebot.TeleBot):
    @bot.message_handler(commands=["start"])
    def start_handler(message: Message):
        user = message.from_user
        if user.username:
            text = f"Assalomu alaykum hurmatli foydalanuvchi @{user.username}\nTugmalardan birini tanlang 👇"
        else:
            text = f"Assalomu alaykum hurmatli foydalanuvchi {user.full_name}\nTugmalardan birini tanlang 👇"

        clear_state(message.chat.id)
        bot.send_message(message.chat.id, text, reply_markup=get_menu())

    @bot.message_handler(func=lambda m: m.text == "Pdf 📁")
    def pdf_button_handler(message: Message):
        set_state(message.chat.id, State.WAITING_PDF)
        bot.send_message(
            message.chat.id,
            "Pdf file tashlang (.pdf)",
            reply_markup=get_menu()
        )

    @bot.message_handler(content_types=["document"], func=lambda m: get_state(m.chat.id) == State.WAITING_PDF)
    def get_pdf_file(message: Message):
        document = message.document

        if not document.file_name or not document.file_name.lower().endswith(".pdf"):
            bot.send_message(message.chat.id, "Faqat .pdf formatdagi fayl yuboring.")
            return

        file_path = make_file_path(PDF_DIR, ".pdf")

        try:
            file_info = bot.get_file(document.file_id)
            downloaded_file = bot.download_file(file_info.file_path)

            with open(file_path, "wb") as new_file:
                new_file.write(downloaded_file)

        except Exception as e:
            bot.send_message(message.chat.id, f"Fayl yuklashda xatolik: {e}")
            return

        user_data[message.chat.id] = {
            "pdf_path": str(file_path),
            "pdf_name": document.file_name
        }

        set_state(message.chat.id, State.WAITING_PDF_CONFIRM)

        bot.send_message(
            message.chat.id,
            "PDFni Wordga o'girmoqchimisiz?",
            reply_markup=get_confirm()
        )

    @bot.message_handler(func=lambda m: get_state(m.chat.id) == State.WAITING_PDF and m.content_type != "document")
    def waiting_pdf_invalid(message: Message):
        bot.send_message(message.chat.id, "Iltimos, .pdf fayl yuboring.")

    @bot.message_handler(func=lambda m: get_state(m.chat.id) == State.WAITING_PDF_CONFIRM and m.text == "Ha ✅")
    def pdf_confirm_yes(message: Message):
        data = user_data.get(message.chat.id, {})
        pdf_path = data.get("pdf_path")
        pdf_name = data.get("pdf_name", "converted.pdf")

        if not pdf_path or not os.path.exists(pdf_path):
            clear_state(message.chat.id)
            bot.send_message(
                message.chat.id,
                "PDF topilmadi. Qaytadan yuboring.",
                reply_markup=get_menu()
            )
            return

        output_docx = str(make_file_path(OUT_DIR, ".docx"))
        progress_msg = bot.send_message(message.chat.id, "Jarayonda... 0%")

        try:
            bot.edit_message_text(
                "Jarayonda... 20%",
                chat_id=message.chat.id,
                message_id=progress_msg.message_id
            )

            extracted_text = extract_text_from_pdf_full(pdf_path)

            if not extracted_text or not extracted_text.strip():
                bot.edit_message_text(
                    "Jarayonda... 100%",
                    chat_id=message.chat.id,
                    message_id=progress_msg.message_id
                )
                bot.send_message(
                    message.chat.id,
                    "PDF ichidan matn topilmadi. Word fayl yaratilmadi.",
                    reply_markup=get_menu()
                )
                return

            bot.edit_message_text(
                "Jarayonda... 60%",
                chat_id=message.chat.id,
                message_id=progress_msg.message_id
            )

            create_docx_from_text(extracted_text, output_docx)

            if not os.path.exists(output_docx):
                raise Exception("Word fayl yaratilmadi.")

            bot.edit_message_text(
                "Jarayonda... 90%",
                chat_id=message.chat.id,
                message_id=progress_msg.message_id
            )

            with open(output_docx, "rb") as f:
                safe_name = f"{Path(pdf_name).stem}.docx"
                bot.send_document(
                    message.chat.id,
                    f,
                    visible_file_name=safe_name
                )

            bot.edit_message_text(
                "Jarayonda... 100% ✅\nWord fayl yuborildi.",
                chat_id=message.chat.id,
                message_id=progress_msg.message_id
            )

            bot.send_message(
                message.chat.id,
                "Orqaga qaytdingiz 🔙",
                reply_markup=get_menu()
            )

        except Exception as e:
            try:
                bot.edit_message_text(
                    "Jarayonda xatolik yuz berdi ❌",
                    chat_id=message.chat.id,
                    message_id=progress_msg.message_id
                )
            except:
                pass

            bot.send_message(message.chat.id, f"Xatolik yuz berdi: {e}")

        finally:
            remove_file(pdf_path)
            remove_file(output_docx)
            user_data.pop(message.chat.id, None)
            clear_state(message.chat.id)


    @bot.message_handler(func=lambda m: get_state(m.chat.id) == State.WAITING_PDF_CONFIRM and m.text == "Yo'q ❌")
    def pdf_confirm_no(message: Message):
        data = user_data.get(message.chat.id, {})
        remove_file(data.get("pdf_path"))
        user_data.pop(message.chat.id, None)
        clear_state(message.chat.id)
        bot.send_message(message.chat.id, "Orqaga qaytdingiz 🔙", reply_markup=get_menu())

    @bot.message_handler(func=lambda m: m.text == "Word 📁")
    def word_handler(message: Message):
        set_state(message.chat.id, State.WAITING_DOCX)
        bot.send_message(message.chat.id, "Word file tashlang (.docx)", reply_markup=get_menu())

    @bot.message_handler(content_types=["document"], func=lambda m: get_state(m.chat.id) == State.WAITING_DOCX)
    def get_docx_file(message: Message):
        document = message.document

        if not document.file_name or not document.file_name.lower().endswith(".docx"):
            bot.send_message(message.chat.id, "Faqat .docx formatdagi fayl yuboring.")
            return

        file_path = make_file_path(DOCX_DIR, ".docx")

        try:
            file_info = bot.get_file(document.file_id)
            downloaded_file = bot.download_file(file_info.file_path)

            with open(file_path, "wb") as new_file:
                new_file.write(downloaded_file)

        except Exception as e:
            bot.send_message(message.chat.id, f"Fayl yuklashda xatolik: {e}")
            return

        user_data[message.chat.id] = {
            "docx_path": str(file_path),
            "docx_name": document.file_name
        }

        set_state(message.chat.id, State.WAITING_DOCX_CONFIRM)

        bot.send_message(
            message.chat.id,
            "Wordni PDFga o'girmoqchimisiz?",
            reply_markup=get_confirm()
        )

    @bot.message_handler(func=lambda m: get_state(m.chat.id) == State.WAITING_DOCX_CONFIRM and m.text == "Ha ✅")
    def docx_confirm_yes(message: Message):
        data = user_data.get(message.chat.id, {})
        docx_path = data.get("docx_path")
        docx_name = data.get("docx_name", "converted.docx")

        if not docx_path or not os.path.exists(docx_path):
            clear_state(message.chat.id)
            bot.send_message(
                message.chat.id,
                "Word fayl topilmadi. Qaytadan yuboring.",
                reply_markup=get_menu()
            )
            return

        output_pdf = None
        progress_msg = bot.send_message(message.chat.id, "Jarayonda... 0%")

        try:
            bot.edit_message_text(
                "Jarayonda... 20%\nWord fayl tekshirilmoqda...",
                chat_id=message.chat.id,
                message_id=progress_msg.message_id
            )

            bot.edit_message_text(
                "Jarayonda... 50%\nPDF ga o'girilmoqda...",
                chat_id=message.chat.id,
                message_id=progress_msg.message_id
            )

            output_pdf = convert_docx_to_pdf_ilovepdf(str(docx_path), str(OUT_DIR))

            if not output_pdf or not os.path.exists(output_pdf):
                raise Exception("PDF fayl yaratilmadi.")

            bot.edit_message_text(
                "Jarayonda... 90%\nPDF fayl yuborilmoqda...",
                chat_id=message.chat.id,
                message_id=progress_msg.message_id
            )

            with open(output_pdf, "rb") as f:
                safe_name = f"{Path(docx_name).stem}.pdf"
                bot.send_document(
                    message.chat.id,
                    f,
                    visible_file_name=safe_name
                )

            bot.edit_message_text(
                "Jarayonda... 100% ✅\nTayyor. PDF yuborildi.",
                chat_id=message.chat.id,
                message_id=progress_msg.message_id
            )

            bot.send_message(
                message.chat.id,
                "Orqaga qaytdingiz 🔙",
                reply_markup=get_menu()
            )

        except Exception as e:
            try:
                bot.edit_message_text(
                    "Jarayonda xatolik yuz berdi ❌",
                    chat_id=message.chat.id,
                    message_id=progress_msg.message_id
                )
            except:
                pass

            bot.send_message(message.chat.id, f"Xatolik yuz berdi: {e}")

        finally:
            remove_file(str(docx_path))
            if output_pdf:
                remove_file(str(output_pdf))
            user_data.pop(message.chat.id, None)
            clear_state(message.chat.id)

    @bot.message_handler(func=lambda m: get_state(m.chat.id) == State.WAITING_DOCX_CONFIRM and m.text == "Yo'q ❌")
    def docx_confirm_no(message: Message):
        data = user_data.get(message.chat.id, {})
        remove_file(data.get("docx_path"))
        user_data.pop(message.chat.id, None)
        clear_state(message.chat.id)

        bot.send_message(
            message.chat.id,
            "Orqaga qaytdingiz 🔙",
            reply_markup=get_menu()
        )

    @bot.message_handler(func=lambda m: m.text == "Text ✉️")
    def text_handler(message: Message):
        set_state(message.chat.id, State.WAITING_TEXT)
        bot.send_message(message.chat.id, "Matn yuboring!", reply_markup=get_menu())

    @bot.message_handler(func=lambda m: get_state(m.chat.id) == State.WAITING_TEXT and m.content_type == "text")
    def get_text_and_convert_to_word(message: Message):
        text = message.text.strip()

        if not text:
            bot.send_message(message.chat.id, "Matn yuboring!")
            return

        output_docx = make_file_path(OUT_DIR, ".docx")

        try:
            create_docx_from_text(text, str(output_docx))
            with open(output_docx, "rb") as f:
                bot.send_document(message.chat.id, f, visible_file_name="matn.docx")

        except Exception as e:
            bot.send_message(message.chat.id, f"Xatolik yuz berdi: {e}")

        finally:
            remove_file(str(output_docx))
            clear_state(message.chat.id)
            bot.send_message(message.chat.id, "Orqaga qaytdingiz 🔙", reply_markup=get_menu())

    @bot.message_handler(func=lambda m: m.text == "Tarjima 🌐")
    def tarjima_handler(message: Message):
        set_state(message.chat.id, State.WAITING_TRANSLATE_FILE)
        bot.send_message(
            message.chat.id,
            "Tarjima uchun pdf yoki docx fayl tashlang",
            reply_markup=get_menu()
        )

        @bot.message_handler(
            content_types=["document"],
            func=lambda m: get_state(m.chat.id) == State.WAITING_TRANSLATE_FILE
        )

        def get_translate_file(message: Message):
            document = message.document
            file_name = (document.file_name or "").lower()

            if not (file_name.endswith(".pdf") or file_name.endswith(".docx")):
                bot.send_message(
                    message.chat.id,
                    "Faqat pdf yoki docx formatdagi fayl yuboring."
                )
                return

            suffix = ".pdf" if file_name.endswith(".pdf") else ".docx"
            original_name = document.file_name or f"file{suffix}"

            try:
                file_path = make_file_path(TRANSLATE_DIR, suffix)

                file_info = bot.get_file(document.file_id)
                downloaded_file = bot.download_file(file_info.file_path)

                with open(file_path, "wb") as new_file:
                    new_file.write(downloaded_file)

                user_data[message.chat.id] = {
                    "translate_file_path": str(file_path),
                    "translate_file_name": original_name,
                    "translate_file_type": suffix
                }

                set_state(message.chat.id, State.WAITING_TRANSLATE_LANG)
                bot.send_message(
                    message.chat.id,
                    "Tarjima tilini tanlang",
                    reply_markup=get_lang()
                )

            except Exception as e:
                bot.send_message(message.chat.id, f"Fayl yuklashda xatolik: {e}")

    @bot.message_handler(func=lambda m: get_state(m.chat.id) == State.WAITING_TRANSLATE_LANG and m.text == "Orqaga 🔙")
    def translate_back(message: Message):
        data = user_data.get(message.chat.id, {})
        remove_file(data.get("translate_file_path"))
        clear_state(message.chat.id)
        bot.send_message(message.chat.id, "Orqaga qaytdingiz 🔙", reply_markup=get_menu())


    @bot.message_handler(func=lambda m: get_state(m.chat.id) == State.WAITING_TRANSLATE_LANG)
    def translate_file_by_lang(message: Message):
        data = user_data.get(message.chat.id, {})
        file_path = data.get("translate_file_path")
        file_name = data.get("translate_file_name", "file")
        file_type = data.get("translate_file_type", "")

        lang_map = {
            "🇺🇿 O'zbek": "uz",
            "🇷🇺 Rus": "ru",
            "🇬🇧 English": "en",
            "🇹🇷 Turk": "tr"
        }

        target_lang = lang_map.get(message.text)
        if not target_lang:
            bot.send_message(message.chat.id, "Iltimos, tugmalardan birini tanlang.")
            return

        if not file_path or not os.path.exists(file_path):
            clear_state(message.chat.id)
            bot.send_message(
                message.chat.id,
                "Fayl topilmadi. Qaytadan yuboring.",
                reply_markup=get_menu()
            )
            return

        progress_msg = bot.send_message(message.chat.id, "Jarayonda... 0%")
        output_file = None

        try:
            bot.edit_message_text(
                "Jarayonda... 20%\nMatn ajratilmoqda...",
                message.chat.id,
                progress_msg.message_id
            )

            extracted_text = ""

            if file_type == ".pdf":
                extracted_text = extract_text_from_pdf_full(file_path)
            elif file_type == ".docx":
                extracted_text = extract_text_from_docx_file_full(file_path)

            if not extracted_text or not extracted_text.strip():
                bot.send_message(message.chat.id, "Fayldan matn ajratilmadi.")
                return

            bot.edit_message_text(
                "Jarayonda... 60%\nTarjima qilinmoqda...",
                message.chat.id,
                progress_msg.message_id
            )

            translated_text = smart_translate_full(extracted_text, target_lang)

            if not translated_text or not translated_text.strip():
                bot.send_message(message.chat.id, "Tarjima natijasi bo‘sh chiqdi.")
                return

            bot.edit_message_text(
                "Jarayonda... 85%\nWord fayl tayyorlanmoqda...",
                message.chat.id,
                progress_msg.message_id
            )

            output_file = str(make_file_path(OUT_DIR, ".docx"))
            create_docx_from_text(translated_text, output_file)

            bot.edit_message_text(
                "Jarayonda... 100% ✅",
                message.chat.id,
                progress_msg.message_id
            )

            with open(output_file, "rb") as f:
                safe_name = f"translated_{Path(file_name).stem}.docx"
                bot.send_document(message.chat.id, f, visible_file_name=safe_name)

        except Exception as e:
            bot.send_message(message.chat.id, f"Xatolik yuz berdi: {e}")

        finally:
            remove_file(file_path)
            if output_file:
                remove_file(output_file)
            clear_state(message.chat.id)
            bot.send_message(message.chat.id, "Orqaga qaytdingiz 🔙", reply_markup=get_menu())

    @bot.message_handler(func=lambda m: m.text == "Adminga murojaat 👨‍💻")
    def admin_handler(message: Message):
        clear_state(message.chat.id)
        bot.send_message(message.chat.id, "Lichka: @musulmon_0319\nTelefon raqam: +998712000540")